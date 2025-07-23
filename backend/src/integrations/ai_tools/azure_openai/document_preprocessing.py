"""
Document Preprocessing Utilities.

This module provides functions for preprocessing various document types
before sending them to Azure OpenAI for analysis.
"""

import os
import logging
import json
from typing import Dict, List, Optional, Any, Union, Tuple
from pathlib import Path
import tempfile

logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except (ImportError, ValueError):
    PANDAS_AVAILABLE = False
    logger.warning("Pandas library not available or incompatible. Excel processing functionality will be limited.")

try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 library not available. PDF processing functionality will be limited.")

try:
    import pdf2image
    PDF_IMAGE_AVAILABLE = True
except ImportError:
    PDF_IMAGE_AVAILABLE = False
    logger.warning("pdf2image library not available. PDF image processing functionality will be limited.")

try:
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logger.warning("openpyxl library not available. Excel processing functionality will be limited.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx library not available. Word document processing functionality will be limited.")

logger = logging.getLogger(__name__)


async def preprocess_document(document_path: str) -> Tuple[str, Optional[str]]:
    """
    Preprocess document based on its file type.
    
    Args:
        document_path: Path to the document file
        
    Returns:
        Tuple containing (document_content, error_message)
    """
    try:
        file_extension = Path(document_path).suffix.lower()
        
        # Process based on file type
        if file_extension == '.pdf':
            return await preprocess_pdf(document_path)
        elif file_extension in ['.xlsx', '.xls']:
            return await preprocess_excel(document_path)
        elif file_extension == '.docx':
            return await preprocess_docx(document_path)
        elif file_extension in ['.json']:
            return await preprocess_json(document_path)
        elif file_extension in ['.txt', '.csv', '.md']:
            return await preprocess_text(document_path)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            # Images would require OCR processing
            return "", "Image preprocessing requires OCR which is not implemented"
        else:
            return "", f"Unsupported file type: {file_extension}"
    
    except Exception as e:
        logger.error(f"Error preprocessing document {document_path}: {str(e)}")
        return "", f"Preprocessing error: {str(e)}"


async def preprocess_pdf(pdf_path: str) -> Tuple[str, Optional[str]]:
    """
    Extract text from PDF document.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Tuple containing (document_content, error_message)
    """
    try:
        with open(pdf_path, 'rb') as file:
            # Create a PDF reader object
            pdf = PdfReader(file)
            
            # Initialize content
            content = []
            
            # Extract text from each page
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text.strip():
                    content.append(f"--- Page {page_num + 1} ---")
                    content.append(text)
            
            # If no text was extracted, the PDF might be scanned
            if not content:
                return "", "PDF appears to be scanned or contains no extractable text"
                
            return "\n".join(content), None
            
    except Exception as e:
        logger.error(f"Error processing PDF {pdf_path}: {str(e)}")
        return "", f"PDF processing error: {str(e)}"


async def preprocess_excel(excel_path: str) -> Tuple[str, Optional[str]]:
    """
    Convert Excel to structured JSON representation.
    
    Args:
        excel_path: Path to the Excel file
        
    Returns:
        Tuple containing (document_content, error_message)
    """
    try:
        # Load Excel file
        result = {"sheets": {}}
        
        # Use pandas to read Excel
        xlsx = pd.ExcelFile(excel_path)
        sheet_names = xlsx.sheet_names
        
        for sheet_name in sheet_names:
            # Read sheet data
            df = pd.read_excel(excel_path, sheet_name=sheet_name)
            
            # Convert to JSON string (orient='records' creates a list of row objects)
            sheet_data = df.to_json(orient='records', date_format='iso')
            
            # Parse back to Python object and add to result
            result["sheets"][sheet_name] = json.loads(sheet_data)
        
        # Convert entire result to formatted JSON string
        return json.dumps(result, indent=2), None
        
    except Exception as e:
        logger.error(f"Error processing Excel {excel_path}: {str(e)}")
        return "", f"Excel processing error: {str(e)}"


async def preprocess_docx(docx_path: str) -> Tuple[str, Optional[str]]:
    """
    Extract text from Word document.
    
    Args:
        docx_path: Path to the Word file
        
    Returns:
        Tuple containing (document_content, error_message)
    """
    try:
        # Open the document
        doc = Document(docx_path)
        
        # Extract text from paragraphs
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content), None
        
    except Exception as e:
        logger.error(f"Error processing Word document {docx_path}: {str(e)}")
        return "", f"Word document processing error: {str(e)}"


async def preprocess_json(json_path: str) -> Tuple[str, Optional[str]]:
    """
    Read and format JSON file.
    
    Args:
        json_path: Path to the JSON file
        
    Returns:
        Tuple containing (document_content, error_message)
    """
    try:
        with open(json_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            
        # Format JSON with indentation
        formatted_json = json.dumps(data, indent=2)
        return formatted_json, None
        
    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in {json_path}: {str(e)}")
        return "", f"Invalid JSON: {str(e)}"
    except Exception as e:
        logger.error(f"Error processing JSON {json_path}: {str(e)}")
        return "", f"JSON processing error: {str(e)}"


async def preprocess_text(text_path: str) -> Tuple[str, Optional[str]]:
    """
    Read text file content.
    
    Args:
        text_path: Path to the text file
        
    Returns:
        Tuple containing (document_content, error_message)
    """
    try:
        # Try different encodings if UTF-8 fails
        encodings = ['utf-8', 'latin-1', 'cp1252']
        content = None
        
        for encoding in encodings:
            try:
                with open(text_path, 'r', encoding=encoding) as file:
                    content = file.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            return "", "Could not decode file with any supported encoding"
            
        return content, None
        
    except Exception as e:
        logger.error(f"Error processing text file {text_path}: {str(e)}")
        return "", f"Text file processing error: {str(e)}"